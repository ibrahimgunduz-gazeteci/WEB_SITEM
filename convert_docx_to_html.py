#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX → HTML dönüştürme - 72 makale için statik HTML oluştur
GitHub Pages'de tarayıcıda makale içeriğini göstermek için
"""

import os
import json
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("❌ python-docx gerekli: pip install python-docx")
    exit(1)

def sanitize_filename(filename):
    """Dosya adını HTML dosyası için safe hale getir"""
    # .docx'i kaldır
    if filename.endswith('.docx'):
        filename = filename[:-5]
    # Spec olmayan karakterleri kaldır
    filename = re.sub(r'[^\w\s-]', '', filename, flags=re.UNICODE)
    # Boşlukları - ile değiştir
    filename = re.sub(r'[\s]+', '-', filename)
    return filename.lower()

def docx_to_html(docx_path):
    """
    .docx dosyasını HTML'e dönüştür (paragraflar, listeler, tablolar, resimler)
    """
    try:
        doc = Document(docx_path)
        html_parts = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # Paragraf formatting
            para_style = para.style.name if para.style else "Normal"
            
            # Heading'ler
            if "Heading" in para_style:
                level = int(para_style[-1]) if para_style[-1].isdigit() else 1
                html_parts.append(f"<h{level}>{para.text}</h{level}>")
                continue
            
            # Normal metin ile formatting
            para_html = []
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                
                # HTML escape
                text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                
                # Formatting
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                if run.underline:
                    text = f"<u>{text}</u>"
                
                para_html.append(text)
            
            if para_html:
                html_parts.append(f"<p>{''.join(para_html)}</p>")
        
        # Tablolar
        for table in doc.tables:
            html_parts.append('<table border="1" style="width:100%; border-collapse:collapse;">')
            for row in table.rows:
                html_parts.append("<tr>")
                for cell in row.cells:
                    html_parts.append(f"<td style='padding:8px; border:1px solid #ddd;'>{cell.text}</td>")
                html_parts.append("</tr>")
            html_parts.append("</table>")
        
        # Resimler (varsa)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                html_parts.append(f"<p><em>[Resim: {rel.target_ref}]</em></p>")
        
        return "\n".join(html_parts)
    
    except Exception as e:
        print(f"❌ Hata {docx_path} dönüştürülürken: {e}")
        return None

def main():
    docs_folder = "assets/documents"
    articles_file = "assets/yazilar.json"
    html_output_folder = "assets/articles-html"
    
    # HTML output klasörü oluştur
    Path(html_output_folder).mkdir(parents=True, exist_ok=True)
    
    # articles.json'u oku
    with open(articles_file, 'r', encoding='utf-8') as f:
        articles = json.load(f)
    
    converted_count = 0
    failed_count = 0
    
    for article in articles:
        docx_filename = article.get('fileName', '')
        docx_path = os.path.join(docs_folder, docx_filename)
        
        if not os.path.exists(docx_path):
            print(f"⚠️  Dosya bulunamadı: {docx_filename}")
            article['htmlFile'] = None
            article['htmlContent'] = None
            failed_count += 1
            continue
        
        print(f"⏳ Dönüştürülüyor: {docx_filename}")
        
        html_content = docx_to_html(docx_path)
        
        if html_content is None:
            print(f"❌ Dönüştürme başarısız: {docx_filename}")
            article['htmlFile'] = None
            article['htmlContent'] = None
            failed_count += 1
            continue
        
        # HTML dosyasını kaydet
        safe_filename = sanitize_filename(docx_filename)
        html_filename = f"{safe_filename}.html"
        html_path = os.path.join(html_output_folder, html_filename)
        
        try:
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            article['htmlFile'] = f"assets/articles-html/{html_filename}"
            article['htmlContent'] = html_content
            converted_count += 1
            print(f"✓ {html_filename}")
        
        except Exception as e:
            print(f"❌ Kaydetme hatası ({html_filename}): {e}")
            article['htmlFile'] = None
            article['htmlContent'] = None
            failed_count += 1
    
    # Güncellenmiş articles.json'u kaydet
    with open(articles_file, 'w', encoding='utf-8') as f:
        json.dump(articles, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ TAMAMLANDI:")
    print(f"   ✓ Dönüştürülen: {converted_count}")
    print(f"   ✗ Başarısız: {failed_count}")
    print(f"   📁 Çıktı: {html_output_folder}/")
    print(f"   📄 Güncellenen: {articles_file}")

if __name__ == "__main__":
    main()
