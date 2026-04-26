#!/usr/bin/env python3
"""
İbrahim Gündüz Websitesi - Belge İşleyici
==========================================
Bu script assets/documents/ klasöründeki .docx dosyalarını tarar ve
assets/yazilar.json dosyasını günceller.

- Mevcut makaleler (zaten JSON'da olanlar) değiştirilmez.
- Sadece YENİ .docx dosyaları işlenir ve JSON'a eklenir.
- Görüntüler assets/images/articles/ klasörüne çıkarılır.

Kullanım:
    python scripts/process_docs.py

Yeni makale eklemek için:
    1. .docx dosyasını assets/documents/ klasörüne koy
    2. GitHub'a push et
    3. GitHub Actions otomatik olarak bu scripti çalıştırır
"""

import os
import sys
import json
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("HATA: python-docx yüklü değil. Çalıştır: pip install python-docx")
    sys.exit(1)

# Yollar (repo kökünden çalıştırılmalı)
DOCS_DIR    = Path("assets/documents")
IMAGES_DIR  = Path("assets/images/articles")
OUTPUT_JSON = Path("assets/yazilar.json")
DEFAULT_THUMBNAIL = "static/ibrahimgunduz.png"

# Türkçe → ASCII dönüşüm tablosu (dosya adı ve görsel adlandırma için)
TR_TABLE = str.maketrans(
    "ğĞüÜşŞıİöÖçÇ",
    "gGuUsSiIoOcC"
)


def transliterate(text: str) -> str:
    """Türkçe karakterleri ASCII karşılıklarıyla değiştirir."""
    return text.translate(TR_TABLE)


def extract_images(doc, base_name: str) -> tuple[dict, list]:
    """
    Docx dosyasındaki tüm resimleri IMAGES_DIR'e kaydeder.
    Dönüş: (rId -> dosya_yolu sözlüğü, sıralı dosya yolları listesi)
    """
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    image_map = {}   # relationship ID -> 'assets/images/articles/...'
    image_list = []  # sıralı liste

    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
            ct   = rel.target_part.content_type

            if "jpeg" in ct or "jpg" in ct:
                ext = "jpg"
            elif "png" in ct:
                ext = "png"
            elif "gif" in ct:
                ext = "gif"
            elif "webp" in ct:
                ext = "webp"
            else:
                ext = "png"

            img_num      = len(image_list) + 1
            safe_base    = transliterate(base_name)
            img_filename = f"{safe_base}-img{img_num}.{ext}"
            img_path     = IMAGES_DIR / img_filename
            rel_path     = f"assets/images/articles/{img_filename}"

            # Dosya zaten varsa üzerine yazma
            if not img_path.exists():
                with open(img_path, "wb") as f:
                    f.write(blob)
                print(f"    Görsel kaydedildi: {img_filename}")

            image_map[rel_id] = rel_path
            image_list.append(rel_path)

        except Exception as e:
            print(f"    UYARI: Görsel çıkarılamadı ({rel_id}): {e}", file=sys.stderr)

    return image_map, image_list


def paragraph_to_html(para, image_map: dict) -> str:
    """Bir docx paragrafını HTML satırına dönüştürür."""
    # Satır içi resim var mı kontrol et
    A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    blips = para._p.findall(f".//{{{A_NS}}}blip")
    if blips:
        parts = []
        for blip in blips:
            r_embed = blip.get(f"{{{R_NS}}}embed")
            if r_embed and r_embed in image_map:
                parts.append(f'<img src="{image_map[r_embed]}" alt="Görsel">')
        if parts:
            text = para.text.strip()
            result = f"<p>{text}</p>\n" if text else ""
            result += "\n".join(parts)
            return result

    text = para.text.strip()
    if not text:
        return ""

    style = para.style.name if para.style else ""
    if "Heading 1" in style or style == "Title":
        return f"<h1>{text}</h1>"
    elif "Heading 2" in style:
        return f"<h2>{text}</h2>"
    elif "Heading" in style:
        return f"<h2>{text}</h2>"
    else:
        return f"<p>{text}</p>"


def process_docx(docx_path: Path) -> dict | None:
    """
    Tek bir .docx dosyasını işler.
    Dönüş: yazilar.json formatında sözlük, ya da hata varsa None.
    """
    filename = docx_path.name
    title    = filename.replace(".docx", "")

    print(f"  İşleniyor: {filename}")

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        print(f"  HATA: Dosya okunamadı: {e}", file=sys.stderr)
        return None

    # Görselleri çıkar
    image_map, image_list = extract_images(doc, title)

    # HTML içeriği oluştur (paragraflar)
    html_parts = []
    for para in doc.paragraphs:
        html = paragraph_to_html(para, image_map)
        if html:
            html_parts.append(html)

    # Tablo varsa ekle
    for table in doc.tables:
        rows_html = []
        for row in table.rows:
            cells = "".join(
                f'<td style="padding:8px; border:1px solid #ddd;">{cell.text.strip()}</td>'
                for cell in row.cells
            )
            rows_html.append(f"<tr>{cells}</tr>")
        table_html = (
            '<table style="border-collapse:collapse; width:100%; margin:20px 0;">'
            + "".join(rows_html)
            + "</table>"
        )
        html_parts.append(table_html)

    # JSON'da görsel referansları yoksa görselleri en sona ekle
    has_inline_images = any('<img' in p for p in html_parts)
    if not has_inline_images and image_list:
        for img_path in image_list:
            html_parts.append(f'<img src="{img_path}" alt="Görsel">')

    html_content = "\n".join(html_parts) if html_parts else f"<p>{title}</p>"
    thumbnail    = image_list[0] if image_list else DEFAULT_THUMBNAIL

    return {
        "title":       title,
        "fileName":    filename,
        "thumbnail":   thumbnail,
        "htmlContent": html_content,
    }


def main():
    # Repo kökünden çalıştırılıp çalıştırılmadığını kontrol et
    if not DOCS_DIR.exists():
        print(f"HATA: Belgeler klasörü bulunamadı: {DOCS_DIR}", file=sys.stderr)
        print("Bu scripti repo kökünden çalıştırın: python scripts/process_docs.py")
        sys.exit(1)

    # Mevcut JSON'u yükle
    existing: dict[str, dict] = {}
    if OUTPUT_JSON.exists():
        try:
            with open(OUTPUT_JSON, "r", encoding="utf-8") as f:
                existing_list = json.load(f)
            existing = {item["fileName"]: item for item in existing_list}
            print(f"Mevcut JSON yüklendi: {len(existing)} makale")
        except Exception as e:
            print(f"UYARI: JSON yüklenemedi, sıfırdan başlanıyor: {e}", file=sys.stderr)

    # documents/ klasöründeki tüm .docx dosyalarını tara
    docx_files = sorted(DOCS_DIR.glob("*.docx"))
    print(f"Belgeler klasöründe {len(docx_files)} .docx dosyası bulundu\n")

    all_articles  = []
    new_count     = 0
    docx_filenames = {p.name for p in docx_files}

    # Silinen belgeleri tespit et
    deleted = [name for name in existing if name not in docx_filenames]
    if deleted:
        print(f"SİLİNEN BELGELER ({len(deleted)}):")
        for name in deleted:
            print(f"  - {name}")
        print()

    for docx_path in docx_files:
        filename = docx_path.name
        if filename in existing:
            # Mevcut makale → değiştirme, olduğu gibi al
            all_articles.append(existing[filename])
        else:
            # Yeni makale → işle
            print(f"YENİ BELGE: {filename}")
            entry = process_docx(docx_path)
            if entry:
                all_articles.append(entry)
                new_count += 1
            else:
                print(f"  Atlandı (işlem başarısız)", file=sys.stderr)

    # Değişiklik yoksa JSON'u değiştirme
    if new_count == 0 and not deleted:
        print("\nDeğişiklik yok. yazilar.json güncellenmedi.")
        return

    # Güncellenmiş JSON'u yaz
    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(all_articles, f, ensure_ascii=False, indent=2)

    summary = []
    if new_count:    summary.append(f"{new_count} yeni eklendi")
    if deleted:      summary.append(f"{len(deleted)} silindi")
    print(f"\nyazilar.json güncellendi: {len(all_articles)} makale ({', '.join(summary)})")


if __name__ == "__main__":
    main()
