#!/usr/bin/env python3
"""
Easy Article Addition System for ibrahimgunduz.com.tr
- Converts DOCX to HTML
- Generates JSON entry
- Updates yazilar.json
- Re-embeds data in all HTML files
- Optionally commits to git
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from datetime import datetime
from docx import Document
import hashlib

BASE_DIR = Path(__file__).parent
ASSETS_DIR = BASE_DIR / "assets"
ARTICLES_HTML_DIR = ASSETS_DIR / "articles-html"
YAZILAR_JSON = ASSETS_DIR / "yazilar.json"

def get_user_input():
    """Get article information from user"""
    print("\n" + "="*60)
    print("YENİ YAZIYA HWelcome EKLE SISTEMI")
    print("="*60)
    
    # Get DOCX file
    while True:
        docx_path = input("\n📄 DOCX dosyasının yolunu girin: ").strip()
        if os.path.exists(docx_path) and docx_path.endswith('.docx'):
            break
        print("❌ Dosya bulunamadı veya .docx değil. Tekrar deneyin.")
    
    # Get article title
    title = input("\n📝 Yazı başlığı: ").strip()
    if not title:
        doc = Document(docx_path)
        title = doc.paragraphs[0].text if doc.paragraphs else "Başlıksız"
    
    # Get thumbnail image path (optional)
    thumbnail = None
    thumb_input = input("\n🖼️  Thumbnail resim yolu (opsiyonel, Enter geç): ").strip()
    if thumb_input and os.path.exists(thumb_input):
        thumbnail = thumb_input
    
    # Get article type
    print("\n📂 Yazı türü seçin:")
    print("1. Makale (article)")
    print("2. Haber (news)")
    article_type = input("Seçim (1-2): ").strip()
    article_type_name = "article" if article_type == "1" else "news"
    
    # Get language
    print("\n🌍 Dil seçin:")
    print("1. Türkçe")
    print("2. İngilizce")
    lang_choice = input("Seçim (1-2): ").strip()
    language = "tr" if lang_choice == "1" else "en"
    
    return {
        'docx_path': docx_path,
        'title': title,
        'thumbnail': thumbnail,
        'type': article_type_name,
        'language': language
    }

def docx_to_html(docx_path):
    """Convert DOCX to HTML"""
    try:
        from docx import Document
        doc = Document(docx_path)
        html_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Convert to HTML based on style
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1]
                    html_content.append(f"<h{level}>{para.text}</h{level}>")
                else:
                    html_content.append(f"<p>{para.text}</p>")
        
        return "\n".join(html_content)
    except Exception as e:
        print(f"❌ DOCX dönüştürme hatası: {e}")
        return None

def generate_filename(title, language="tr"):
    """Generate filename from title"""
    # Remove special chars, convert to lowercase with dash
    safe_name = "".join(c if c.isalnum() or c.isspace() else "" for c in title.lower())
    safe_name = "-".join(safe_name.split())[:50]  # Max 50 chars
    if language == "en":
        safe_name = f"en-{safe_name}"
    return safe_name

def add_article(article_info):
    """Add new article to system"""
    try:
        # Convert DOCX to HTML
        print("\n⏳ DOCX dosyası dönüştürülüyor...")
        html_content = docx_to_html(article_info['docx_path'])
        if not html_content:
            return False
        
        # Generate filename
        filename = generate_filename(article_info['title'], article_info['language'])
        
        # Create JSON entry
        article_entry = {
            "fileName": filename,
            "title": article_info['title'],
            "thumbnail": article_info['thumbnail'] or "default.jpg",
            "content": f"Yeni yazı: {article_info['title']}",
            "htmlContent": html_content,
            "type": article_info['type'],
            "language": article_info['language'],
            "dateAdded": datetime.now().isoformat()
        }
        
        # Load existing articles
        if YAZILAR_JSON.exists():
            with open(YAZILAR_JSON, 'r', encoding='utf-8') as f:
                articles = json.load(f)
        else:
            articles = []
        
        # Check for duplicates
        if any(a['fileName'] == filename for a in articles):
            response = input(f"\n⚠️  '{filename}' zaten var. Değiştir? (e/h): ")
            if response.lower() != 'e':
                return False
            articles = [a for a in articles if a['fileName'] != filename]
        
        # Add new article
        articles.append(article_entry)
        
        # Save updated JSON
        print("💾 yazilar.json güncelleniyor...")
        with open(YAZILAR_JSON, 'w', encoding='utf-8') as f:
            json.dump(articles, f, ensure_ascii=False, indent=2)
        
        # Re-embed in HTML files
        print("🔄 HTML dosyaları güncelleniyor...")
        embed_articles_in_html(articles)
        
        print("\n✅ Yazı başarıyla eklendi!")
        print(f"   Dosya: {filename}")
        print(f"   Başlık: {article_info['title']}")
        print(f"   Toplam yazı: {len(articles)}")
        
        # Ask for git commit
        commit = input("\n📦 Git'e commit yap? (e/h): ")
        if commit.lower() == 'e':
            try:
                os.chdir(BASE_DIR)
                subprocess.run(['git', 'add', '.'], check=True)
                subprocess.run(['git', 'commit', '-m', f"Add article: {article_info['title']}"], check=True)
                print("✅ Git commit başarılı!")
            except Exception as e:
                print(f"⚠️  Git commit başarısız: {e}")
        
        return True
        
    except Exception as e:
        print(f"\n❌ Hata: {e}")
        return False

def embed_articles_in_html(articles):
    """Embed articles in HTML files"""
    # Convert to JSON string
    json_str = json.dumps(articles, ensure_ascii=False, indent=2)
    json_code = f"window.articlesData = {json_str};"
    
    # Files to update
    target_files = [
        BASE_DIR / "index.html",
        BASE_DIR / "tum-yazilar.html",
        BASE_DIR / "haber.html"
    ]
    
    for html_file in target_files:
        if not html_file.exists():
            continue
        
        try:
            with open(html_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Remove old window.articlesData if exists
            start = content.find("<script>window.articlesData")
            if start != -1:
                end = content.find("</script>", start) + 9
                content = content[:start] + content[end:]
            
            # Insert new data
            head_end = content.find("</head>")
            if head_end != -1:
                new_content = content[:head_end] + f"\n  <script>{json_code}</script>\n  " + content[head_end:]
                with open(html_file, 'w', encoding='utf-8') as f:
                    f.write(new_content)
                print(f"   ✓ {html_file.name}")
        except Exception as e:
            print(f"   ❌ {html_file.name}: {e}")

def main():
    print("\n🚀 İbrahim Gündüz - Yazı Ekleme Sistemi")
    print("Basit ve hızlı yazı ekle!")
    
    # Get user input
    article_info = get_user_input()
    
    # Add article
    if add_article(article_info):
        sys.exit(0)
    else:
        print("\n❌ İşlem iptal edildi.")
        sys.exit(1)

if __name__ == "__main__":
    main()
