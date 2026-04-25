#!/usr/bin/env python3
"""
Convert DOCX to HTML with proper image extraction
Uses ZIP extraction since DOCX is a compressed XML/media container
"""

import os
import json
import re
import zipfile
import shutil
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).parent
ASSETS_DIR = BASE_DIR / "assets"
ARTICLES_HTML_DIR = ASSETS_DIR / "articles-html"
IMAGES_DIR = ASSETS_DIR / "images" / "articles"
YAZILAR_JSON = ASSETS_DIR / "yazilar.json"

def sanitize_filename(filename):
    """Convert filename to safe format"""
    # Remove special chars, Turkish chars, keep only alphanumeric and dashes
    safe = re.sub(r'[^a-z0-9\-]', '', filename.lower().replace(' ', '-'))
    return safe[:50]

def extract_images_from_docx(docx_path, article_title):
    """Extract images from DOCX ZIP container"""
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    
    extracted_images = []
    safe_title = sanitize_filename(article_title)
    img_counter = 1
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            # Find all media files in word/media/
            for file_info in zip_ref.filelist:
                if 'word/media/' in file_info.filename:
                    # Extract image
                    img_data = zip_ref.read(file_info.filename)
                    
                    # Determine extension
                    ext = file_info.filename.split('.')[-1].lower()
                    if ext == 'jpeg':
                        ext = 'jpg'
                    
                    # Save with clean name
                    img_filename = f"{safe_title}-img{img_counter}.{ext}"
                    img_path = IMAGES_DIR / img_filename
                    
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                    
                    extracted_images.append({
                        'filename': img_filename,
                        'path': f"assets/images/articles/{img_filename}",
                        'original': file_info.filename
                    })
                    img_counter += 1
    except Exception as e:
        print(f"⚠️  Resim çıkartma hatası: {e}")
    
    return extracted_images

def has_inline_image(para):
    """Check if a paragraph contains an embedded inline image via XML"""
    ns_draw = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    ns_vml = 'urn:schemas-microsoft-com:vml'
    return (
        para._element.find(f'.//{{{ns_draw}}}inline') is not None or
        para._element.find(f'.//{{{ns_draw}}}anchor') is not None or
        para._element.find(f'.//{{{ns_vml}}}shape') is not None
    )

def docx_to_html(docx_path, article_title):
    """Convert DOCX to HTML with images placed inline at their original positions"""
    try:
        # Extract images
        images = extract_images_from_docx(docx_path, article_title)
        
        # Load document
        doc = Document(docx_path)
        html_content = []
        first_image = images[0]['path'] if images else None
        img_index = 0
        
        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # If paragraph contains an inline image, place the image here
            if has_inline_image(para):
                if img_index < len(images):
                    img = images[img_index]
                    # Include any caption text alongside the image (skip [Resim:] placeholders)
                    if text and '[resim:' not in text.lower():
                        html_content.append(f'<p class="foto-caption">{text}</p>')
                    html_content.append(f'<figure class="article-image"><img src="{img["path"]}" alt="Makale görseli" class="article-img"/></figure>')
                    img_index += 1
                continue
            
            # Skip empty and old-style image reference text
            if not text or '[resim:' in text.lower():
                continue
            
            # Add as paragraph
            if para.style.name.startswith('Heading'):
                level = para.style.name[-1] if para.style.name[-1].isdigit() else '2'
                html_content.append(f"<h{level}>{text}</h{level}>")
            else:
                html_content.append(f"<p>{text}</p>")
        
        # Append any remaining images that weren't matched to inline positions
        for img in images[img_index:]:
            html_content.append(f'<figure class="article-image"><img src="{img["path"]}" alt="Makale görseli" class="article-img"/></figure>')
        
        return "\n".join(html_content), first_image
        
    except Exception as e:
        print(f"❌ Dönüştürme hatası: {e}")
        return None, None

def process_all_docx():
    """Process all DOCX files in workspace"""
    ARTICLES_HTML_DIR.mkdir(parents=True, exist_ok=True)
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    
    # Find DOCX files in root and subdirectories
    docx_files = list(BASE_DIR.glob("*.docx")) + list(BASE_DIR.glob("**/*.docx"))
    docx_files = list(set(docx_files))  # Remove duplicates
    
    # Load existing articles
    if YAZILAR_JSON.exists():
        with open(YAZILAR_JSON, 'r', encoding='utf-8') as f:
            articles = json.load(f)
    else:
        articles = []
    
    success = 0
    failed = 0
    
    print(f"\n📚 {len(docx_files)} DOCX dosyası bulundu\n")
    
    for docx_file in sorted(docx_files):
        if docx_file.name.startswith('~'):
            continue
        
        print(f"🔄 {docx_file.name}...")
        
        try:
            doc = Document(docx_file)
            title = doc.paragraphs[0].text.strip() if doc.paragraphs else docx_file.stem
            
            # Convert
            html_content, thumbnail = docx_to_html(str(docx_file), title)
            
            if not html_content:
                failed += 1
                continue
            
            # Generate filename
            filename = sanitize_filename(title)
            if not filename:
                filename = docx_file.stem.lower()
            
            html_file = ARTICLES_HTML_DIR / f"{filename}.html"
            
            # Save HTML
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Update article entry
            article = {
                "fileName": docx_file.name,
                "title": title,
                "thumbnail": thumbnail or "assets/images/default-thumbnail.jpg",
                "htmlFile": str(html_file.relative_to(BASE_DIR)),
                "htmlContent": html_content
            }
            
            # Check if exists
            idx = next((i for i, a in enumerate(articles) if a.get('fileName') == title), -1)
            if idx >= 0:
                articles[idx] = article
            else:
                articles.append(article)
            
            print(f"   ✓ {filename}.html ({len(html_content)} byte)")
            success += 1
            
        except Exception as e:
            print(f"   ✗ Hata: {str(e)[:50]}")
            failed += 1
    
    # Save JSON
    with open(YAZILAR_JSON, 'w', encoding='utf-8') as f:
        json.dump(articles, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ Tamamlandı: {success} ✓, {failed} ✗")
    print(f"📄 Toplam {len(articles)} makale")

if __name__ == "__main__":
    process_all_docx()
